import json
import os
import boto3
from opensearchpy import OpenSearch, RequestsHttpConnection
from requests_aws4auth import AWS4Auth
from docx import Document
from io import BytesIO

def get_aws_auth():
    """Create AWS authentication for OpenSearch"""
    credentials = boto3.Session().get_credentials()
    region = os.environ['AWS_REGION']
    return AWS4Auth(
        credentials.access_key,
        credentials.secret_key,
        region,
        'es',
        session_token=credentials.token
    )

def create_opensearch_client():
    """Create OpenSearch client"""
    host = os.environ['OPENSEARCH_ENDPOINT']
    aws_auth = get_aws_auth()
    
    return OpenSearch(
        hosts=[{'host': host, 'port': 443}],
        http_auth=aws_auth,
        use_ssl=True,
        verify_certs=True,
        connection_class=RequestsHttpConnection
    )

def extract_text_from_docx(file_content):
    """Extract text from DOCX file"""
    doc = Document(BytesIO(file_content))
    
    # Extract basic document information
    properties = doc.core_properties
    
    # Extract text from paragraphs
    full_text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Only add non-empty paragraphs
            full_text.append(paragraph.text.strip())
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():  # Only add non-empty cells
                    full_text.append(cell.text.strip())
    
    # Create document metadata
    metadata = {
        'author': str(properties.author) if properties.author else None,
        'created': str(properties.created) if properties.created else None,
        'last_modified': str(properties.modified) if properties.modified else None,
        'title': str(properties.title) if properties.title else None,
        'subject': str(properties.subject) if properties.subject else None,
    }
    
    return {
        'content': ' '.join(full_text),
        'metadata': metadata
    }

def process_file(s3_client, bucket, key):
    """Process the DOCX file from S3"""
    try:
        # Get the file from S3
        response = s3_client.get_object(Bucket=bucket, Key=key)
        file_content = response['Body'].read()
        
        # Extract text and metadata
        document_data = extract_text_from_docx(file_content)
        
        # Add S3 location information
        document_data['source'] = {
            'bucket': bucket,
            'key': key,
            'last_modified': str(response['LastModified']),
            'size': response['ContentLength']
        }
        
        return document_data
    except Exception as e:
        print(f"Error processing file {key}: {str(e)}")
        raise

def create_index_if_not_exists(client, index_name):
    """Create index with mapping if it doesn't exist"""
    try:
        if not client.indices.exists(index=index_name):
            mapping = {
                "mappings": {
                    "properties": {
                        "content": {
                            "type": "text",
                            "analyzer": "standard"
                        },
                        "metadata": {
                            "properties": {
                                "author": {"type": "keyword"},
                                "created": {"type": "date", "ignore_malformed": True},
                                "last_modified": {"type": "date", "ignore_malformed": True},
                                "title": {"type": "text"},
                                "subject": {"type": "text"}
                            }
                        },
                        "source": {
                            "properties": {
                                "bucket": {"type": "keyword"},
                                "key": {"type": "keyword"},
                                "last_modified": {"type": "date", "ignore_malformed": True},
                                "size": {"type": "long"}
                            }
                        }
                    }
                }
            }
            client.indices.create(index=index_name, body=mapping)
            print(f"Created index {index_name}")
    except Exception as e:
        print(f"Error creating index: {str(e)}")
        raise

def handler(event, context):
    """Lambda handler"""
    try:
        s3_client = boto3.client('s3')
        opensearch_client = create_opensearch_client()
        index_name = os.environ.get('OPENSEARCH_INDEX', 'documents')
        
        # Ensure index exists
        create_index_if_not_exists(opensearch_client, index_name)
        
        # Process each record
        processed_records = []
        failed_records = []
        
        for record in event['Records']:
            try:
                # Parse SQS message to get S3 event
                body = json.loads(record['body'])
                s3_record = body['Records'][0]
                bucket = s3_record['s3']['bucket']['name']
                key = s3_record['s3']['object']['key']
                
                # Only process .docx files
                if not key.lower().endswith('.docx'):
                    print(f"Skipping non-DOCX file: {key}")
                    continue
                
                # Process the file
                document_data = process_file(s3_client, bucket, key)
                
                # Index the document
                response = opensearch_client.index(
                    index=index_name,
                    body=document_data,
                    id=f"{bucket}/{key}",  # Use S3 location as document ID
                    refresh=True
                )
                
                print(f"Successfully indexed document {key}: {response}")
                processed_records.append(key)
                
            except Exception as e:
                print(f"Error processing record: {str(e)}")
                failed_records.append(key)
                continue
        
        return {
            'statusCode': 200,
            'body': json.dumps({
                'processed': processed_records,
                'failed': failed_records
            })
        }
        
    except Exception as e:
        print(f"Fatal error: {str(e)}")
        raise